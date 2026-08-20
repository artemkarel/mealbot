"""Импорт плана из любого документа: текст извлекаем сами, структуру разбирает Claude.

PDF — pypdf, Word — python-docx, Excel — openpyxl, остальное читаем как текст.
Claude возвращает строгий JSON той же структуры, что и обычные планы.
"""
import io, json


PARSE_SYSTEM = (
    "Ты разбираешь план питания из документа в строгий JSON. Верни ТОЛЬКО JSON, "
    "без пояснений и без markdown. Схема:\n"
    '{"title": "короткое название плана", "days": [{"day": "Понедельник", '
    '"meals": [{"meal": "Завтрак", "optional": false, "note": null, '
    '"items": [{"name": "Овсяная каша", "qty_min": 180, "qty_max": 200, "unit": "г"}], '
    '"links": [{"name": "Котлеты из индейки «Диетические»", "url": "https://…"}]}]}]}\n'
    "Правила: дни — полные русские названия по порядку с понедельника, максимум 7 "
    "(если в документе дни не названы — считай их подряд); приёмы пищи только из "
    "списка: Завтрак, Второй завтрак, Обед, Полдник, Ужин, Второй ужин "
    "(перекус между завтраком и обедом — Второй завтрак, иначе Полдник); "
    "optional=true только если приём помечен «по желанию»; количества числами в "
    "г/мл/шт: «180–200 г» → qty_min 180 и qty_max 200, одно число → qty_min, "
    "без количества → null; note — только полезное примечание к приёму, если его "
    "нет — null; названия блюд бери из документа, коротко и без количеств; "
    "links — ссылки на товары из примечаний приёма: name — название товара из подписи "
    "перед ссылкой (без лишних слов), url — сам адрес; если ссылок нет — пустой список."
)


def extract_text(filename, data):
    """Достаёт текст из файла по расширению. ValueError — понятное сообщение пользователю."""
    name = (filename or "").lower()
    if name.endswith(".docx"):
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for t in doc.tables:
            for row in t.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells))
        return "\n".join(parts)
    if name.endswith(".pdf"):
        from pypdf import PdfReader
        return "\n".join((p.extract_text() or "")
                         for p in PdfReader(io.BytesIO(data)).pages)
    if name.endswith((".xlsx", ".xlsm")):
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        lines = []
        for ws in wb.worksheets:
            lines.append("== {} ==".format(ws.title))
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines)
    if name.endswith((".doc", ".xls")):
        raise ValueError("старый формат Office не читается — сохрани как .docx/.xlsx или PDF")
    return data.decode("utf-8", "ignore")   # .txt и всё остальное — как текст


IMAGE_TYPES = {"jpg": "image/jpeg", "jpeg": "image/jpeg",
               "png": "image/png", "webp": "image/webp", "gif": "image/gif"}


def image_media_type(filename):
    """image/jpeg и т.п. для фото, None для документов."""
    ext = (filename or "").lower().rsplit(".", 1)[-1]
    return IMAGE_TYPES.get(ext)


def _to_plan(ans):
    s = ans[ans.find("{"):ans.rfind("}") + 1]
    data = json.loads(s)
    days = data.get("days") or []
    if not isinstance(days, list) or not days:
        raise ValueError("в ответе нет дней плана")
    if not any((m.get("items") or []) for d in days for m in (d.get("meals") or [])):
        raise ValueError("в ответе нет ни одного блюда")
    return data


async def parse_plan(text):
    """Текст документа -> {"title", "days": [...]} через Claude."""
    from app.ai import ask_claude
    ans = await ask_claude(PARSE_SYSTEM, [{"role": "user", "content": text[:60000]}],
                           max_tokens=8000, timeout=180)
    return _to_plan(ans)


async def parse_plan_image(data_b64, media_type):
    """Фото плана -> {"title", "days": [...]} через Claude Vision."""
    from app.ai import ask_claude
    content = [{"type": "image",
                "source": {"type": "base64", "media_type": media_type, "data": data_b64}},
               {"type": "text", "text": "Разбери план питания с этого изображения."}]
    ans = await ask_claude(PARSE_SYSTEM, [{"role": "user", "content": content}],
                           max_tokens=8000, timeout=180)
    return _to_plan(ans)
