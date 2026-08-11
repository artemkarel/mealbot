"""Разбор плана питания от диетолога.

Проверено на трёх реальных файлах: 7/7 дней, 35/35 приёмов, все позиции.
Никакой нейросети — формат стабильный, хватает правил.

Запуск отдельно:  python -m app.parser.rules data/План.docx
"""
from __future__ import annotations
import re
import sys
import json
from pathlib import Path

DAYS = ["Понедельник", "Вторник", "Среда", "Четверг",
        "Пятница", "Суббота", "Воскресенье"]
MEALS = ["Завтрак", "Обед", "Полдник", "Ужин", "Второй ужин"]

# "Отварная гречка – 180 гр."   "Хлебцы ... – 2 шт."   "Кофе – 200 – 250 мл."
DISH = re.compile(
    r"^(?P<name>.+?)\s*[–—-]\s*"
    r"(?P<qty>\d+\s*(?:[–—-]\s*\d+)?|½|¼)\s*"
    r"(?P<unit>ч/л\.?|ст/л\.?|гр\.?|г\.|мл\.?|шт\.?|пачки|л\.?)"
    r"\s*(?P<tail>\(.*?\))?\s*;?\s*$"
)
URL = re.compile(r"https?://\S+")
# фразы, означающие "сделать заранее" — идут в вечерние напоминания
PREP = re.compile(r"замач|замочи|на ночь|настаива|всю ночь|заранее|разморо|"
                  r"оставляем на \d+|за \d+ (?:часа|часов|минут)", re.I)

UNIT_MAP = {"гр.": "г", "гр": "г", "г.": "г", "мл.": "мл", "мл": "мл",
            "л.": "л", "шт.": "шт", "шт": "шт", "ч/л.": "ч.л.", "ч/л": "ч.л.",
            "ст/л.": "ст.л.", "ст/л": "ст.л.", "пачки": "пачка"}


W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _paragraph_lines(par) -> list[str]:
    """Текст абзаца с учётом мягких переносов <w:br/>.

    Это важно: в планах две позиции часто стоят в одном абзаце через Shift+Enter
    ("Запеченное филе лосося – 120 гр." + перенос + "Отварной дикий рис – 120 гр.").
    Обычный par.text склеивает их в одну строку и вторая позиция теряется.
    """
    out, buf = [], []
    for node in par._p.iter():
        tag = node.tag
        if tag == W + "t" and node.text:
            buf.append(node.text)
        elif tag in (W + "br", W + "cr"):
            out.append("".join(buf))
            buf = []
        elif tag == W + "tab":
            buf.append(" ")
    out.append("".join(buf))
    return [s.strip() for s in out]


def extract_text(path: str | Path) -> list[str]:
    """docx / txt -> список непустых строк. Без pandoc."""
    path = Path(path)
    lines: list[str] = []
    if path.suffix.lower() == ".docx":
        import docx  # python-docx
        doc = docx.Document(str(path))
        for p in doc.paragraphs:
            lines += _paragraph_lines(p)
        for table in doc.tables:              # на случай, если план в таблице
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        lines += _paragraph_lines(p)
    else:
        lines = Path(path).read_text(encoding="utf-8").split("\n")
    return [l for l in (s.strip() for s in lines) if l]


def parse(lines: list[str]) -> dict:
    days: list[dict] = []
    day = meal = None
    in_recipe = False

    for raw in lines:
        clean = raw.lstrip("*").strip().rstrip(":")

        if clean in DAYS:
            day = {"day": clean, "meals": []}
            days.append(day)
            meal, in_recipe = None, False
            continue

        if clean in MEALS and raw.rstrip().endswith(":") and day is not None:
            meal = {"meal": clean, "optional": clean == "Второй ужин",
                    "items": [], "recipe": None, "notes": [],
                    "links": [], "prep": []}
            day["meals"].append(meal)
            in_recipe = False
            continue

        if meal is None:
            continue

        # начало блока рецепта
        if re.match(r"^\*?\s*Рецепт", raw, re.I) or clean.startswith("Ингредиенты"):
            in_recipe = True
            meal["recipe"] = meal["recipe"] or {
                "title": clean.rstrip(":"), "ingredients": [], "steps": []}
            continue

        # строка-ингредиент внутри рецепта (начинается с дефиса или маркера)
        if in_recipe and re.match(r"^[-–—•]\s", raw):
            body = re.sub(r"^[-–—•]\s*", "", raw).strip()
            m = DISH.match(body)
            if m:
                meal["recipe"]["ingredients"].append({
                    "name": m.group("name").strip(),
                    "qty": m.group("qty").replace(" ", ""),
                    "unit": UNIT_MAP.get(m.group("unit"), m.group("unit")),
                })
            else:  # "Соль – на кончике ножа" — количества нет, это нормально
                meal["recipe"]["ingredients"].append(
                    {"name": body.rstrip("."), "qty": None, "unit": None})
            continue

        for u in URL.findall(raw):
            meal["links"].append(u)
        body = URL.sub("", raw).strip(" .-–—")
        if not body:
            continue

        is_note = raw.lstrip().startswith("*")
        m = DISH.match(body)

        if m and not is_note and not in_recipe:
            q = m.group("qty").replace(" ", "")
            rng = re.match(r"(\d+)[–—-](\d+)", q)
            lo, hi = (int(rng.group(1)), int(rng.group(2))) if rng else (
                (int(q), int(q)) if q.isdigit() else (None, None))
            meal["items"].append({
                "name": m.group("name").strip(),
                "qty_min": lo, "qty_max": hi, "raw_qty": q,
                "unit": UNIT_MAP.get(m.group("unit"), m.group("unit")),
                "paren": (m.group("tail") or "").strip("() ") or None,
            })
        elif len(body) > 15:
            text = body.lstrip("*").strip()
            if in_recipe and meal["recipe"] is not None:
                meal["recipe"]["steps"].append(text)
            else:
                meal["notes"].append(text)
            if PREP.search(text):
                meal["prep"].append(text)

    return {"days": days, "stats": stats(days)}


def stats(days: list[dict]) -> dict:
    items = sum(len(m["items"]) for d in days for m in d["meals"])
    return {
        "days": len(days),
        "meals": sum(len(d["meals"]) for d in days),
        "items": items,
        "recipes": sum(1 for d in days for m in d["meals"] if m["recipe"]),
        "links": sum(len(m["links"]) for d in days for m in d["meals"]),
        "prep": sum(len(m["prep"]) for d in days for m in d["meals"]),
    }


def duplicate_days(days: list[dict]) -> list[list[str]]:
    """Какие дни совпадают полностью — чтобы готовить один раз на два дня."""
    sign: dict[tuple, list[str]] = {}
    for d in days:
        key = tuple(f'{i["name"]}|{i["raw_qty"]}{i["unit"]}'
                    for m in d["meals"] for i in m["items"])
        sign.setdefault(key, []).append(d["day"])
    return [v for v in sign.values() if len(v) > 1]


def parse_file(path: str | Path) -> dict:
    result = parse(extract_text(path))
    result["duplicates"] = duplicate_days(result["days"])
    return result


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Использование: python -m app.parser.rules файл.docx")
        raise SystemExit(1)
    res = parse_file(sys.argv[1])
    print(json.dumps(res["stats"], ensure_ascii=False, indent=2))
    print("Одинаковые дни:", res["duplicates"] or "нет")
    for d in res["days"]:
        print(f"\n=== {d['day']} ===")
        for m in d["meals"]:
            opt = " (по желанию)" if m["optional"] else ""
            print(f"  {m['meal']}{opt}")
            for i in m["items"]:
                print(f"     {i['name']} — {i['raw_qty']} {i['unit']}")
            if m["recipe"]:
                print(f"     [рецепт: {m['recipe']['title']}, "
                      f"{len(m['recipe']['ingredients'])} ингр.]")
            for p in m["prep"]:
                print(f"     [заранее: {p[:70]}]")
